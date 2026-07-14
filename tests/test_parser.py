from pathlib import Path

from docx import Document

from hal_assistant.models import PublicationType
from hal_assistant.parser import (
    disambiguate_generic_title,
    extract_book_title,
    extract_conference_metadata,
    extract_editors,
    extract_publisher_metadata,
    extract_journal_title,
    extract_title,
    normalize_centuries,
    parse_citation,
    parse_docx,
)


def test_extract_quoted_title() -> None:
    assert extract_title("« Un titre », in Une revue, 2024, p.1-10.") == "Un titre"


def test_extract_title_with_nested_french_guillemets() -> None:
    citation = (
        "« Qu’est-ce qu’un « mauvais » théâtre de science-fiction (avant 1920) ? », "
        "in Res Futurae, n°18, 2021."
    )
    assert (
        extract_title(citation)
        == "Qu’est-ce qu’un « mauvais » théâtre de science-fiction (avant 1920) ?"
    )


def test_quoted_title_stops_before_later_quoted_collection() -> None:
    citation = (
        "« Avant-propos : après Kleist, rivalités et artifices », in Florence Fix "
        "(éd.), Jeu d’acteur et corps artificiels, Paris, Orizons, "
        "« comparaisons », 2019, p.9-22."
    )
    assert extract_title(citation) == (
        "Avant-propos : après Kleist, rivalités et artifices"
    )
    assert extract_book_title(citation) == "Jeu d’acteur et corps artificiels"


def test_generic_title_is_qualified_with_french_preposition() -> None:
    assert disambiguate_generic_title("Avant-propos", "Théâtre et science") == (
        "Avant-propos à Théâtre et science"
    )


def test_extract_publisher_and_city_ignores_conference_note() -> None:
    citation = (
        "« Article », [colloque à Paris, 2022], in Livre, "
        "Dijon, EUD, 2024, p.1-10."
    )
    assert extract_publisher_metadata(citation) == ("EUD", "Dijon")
    assert disambiguate_generic_title("Introduction", "Le Mélodramatique") == (
        "Introduction au Mélodramatique"
    )
    assert disambiguate_generic_title("Préface", "Les Années théâtrales") == (
        "Préface aux Années théâtrales"
    )


def test_normalize_centuries_uppercases_explicit_roman_numerals() -> None:
    assert normalize_centuries("Femmes de spectacle au xixe siècle") == (
        "Femmes de spectacle au XIXe siècle"
    )
    assert normalize_centuries("du xviie-xixe siècles") == "du XVIIe-XIXe siècles"


def test_normalize_centuries_does_not_change_ordinary_words() -> None:
    assert normalize_centuries("Pauvreté, maladie et fin de vie") == (
        "Pauvreté, maladie et fin de vie"
    )


def test_normalize_centuries_repairs_known_source_typo() -> None:
    assert normalize_centuries("Growing Old in Nineteenh-Century France") == (
        "Growing Old in Nineteenth-Century France"
    )


def test_extract_book_title_after_editor_prefix() -> None:
    citation = (
        "« La faim après la faim », avant-propos in Florence Fix (éd.), "
        "Manger et être mangé, L’alimentation et ses récits, Paris, Orizons, 2016, p.11-25."
    )
    assert extract_book_title(citation) == (
        "Manger et être mangé, L’alimentation et ses récits"
    )


def test_extract_journal_title_from_article_citation() -> None:
    citation = (
        "« Le corps pétrifié », in Études Francophones, vol. 37, n°2, "
        "Représentations du corps, automne 2024, p.1-17."
    )
    assert extract_journal_title(citation) == "Études Francophones"


def test_extract_journal_title_after_issue_editor() -> None:
    citation = (
        "« Jeanne d’Arc », in Laurence Le Diagon-Jacquin (éd.), "
        "Le Paon d’Héra, N°8, Dijon, 2011, p.155-164."
    )
    assert extract_journal_title(citation) == "Le Paon d’Héra"


def test_extract_conference_title_city_and_country_without_inventing_date() -> None:
    citation = (
        "« Fragmentation », in Peter Schnyder (éd.), De l’écriture et des fragments. "
        "Fragmentation et sciences humaines, [actes du colloque à l’Université de "
        "Haute-Alsace, Mulhouse, 2014], Paris, Classiques Garnier, 2016, p.151-163."
    )
    metadata = extract_conference_metadata(citation)
    assert metadata["conference_title"] == (
        "De l’écriture et des fragments. Fragmentation et sciences humaines"
    )
    assert metadata["conference_city"] == "Mulhouse"
    assert metadata["conference_country"] == "France"
    assert metadata["conference_start_date"] is None
    assert metadata["conference_year_evidence"] == "2014"


def test_extract_conference_container_editors_and_publisher() -> None:
    citation = (
        "« Fragmentation », in Peter Schnyder et Frédérique Toudoire-Surlapierre "
        "(éd.), De l’écriture et des fragments. Fragmentation et sciences humaines, "
        "[actes du colloque à Mulhouse, 2014], Paris, Classiques Garnier, 2016, "
        "p.151-163."
    )
    publication = parse_citation(
        citation,
        "Communication dans un congrès",
        PublicationType.CONFERENCE_PAPER,
        1,
        "Florence Fix",
    )
    assert publication.book_title == (
        "De l’écriture et des fragments. Fragmentation et sciences humaines"
    )
    assert publication.editors == [
        "Peter Schnyder",
        "Frédérique Toudoire-Surlapierre",
    ]
    assert publication.publisher == "Classiques Garnier"
    assert publication.publisher_city == "Paris"


def test_extract_three_explicit_editors() -> None:
    citation = (
        "« Article », in Pierre Watt, Aurore Montesi, Florence Fix (éd.), "
        "Après le temps des rois, Paris, Orizons, 2019, p.61-75."
    )
    assert extract_editors(citation) == [
        "Pierre Watt",
        "Aurore Montesi",
        "Florence Fix",
    ]


def test_editor_extraction_ignores_front_matter_and_conference_notes() -> None:
    assert extract_editors(
        "« Avant-propos », avec Xavier Bonnier, in Xavier Bonnier et Florence "
        "Fix (éd.), Le Détour du comparant, Paris, Classiques Garnier, 2025."
    ) == ["Xavier Bonnier", "Florence Fix"]
    assert extract_editors(
        "« Article » [Actes du colloque, 2024], in Natalia Arregui Barragán et "
        "Carmen Alberdi Urquizu (éd.), L’excès, 2025."
    ) == ["Natalia Arregui Barragán", "Carmen Alberdi Urquizu"]


def test_later_print_version_supplies_container_publisher_place_and_pages() -> None:
    citation = (
        "« Jeux », in Juliette Vion-Dury et Élisabeth Belmas (éd.), Le jeu dans "
        "tous ses états [actes de séminaire], mise en ligne, Paris, CNRS, 2017, "
        "p.1-12. Publication sous le titre Le jeu dans tous ses états. Approches "
        "pluridisciplinaires du phénomène ludique, Paris et Turin, L’Harmattan, "
        "2020, p.212-230."
    )
    publication = parse_citation(
        citation,
        "Communication dans un congrès",
        PublicationType.CONFERENCE_PAPER,
        1,
        "Florence Fix",
    )
    assert publication.year == 2020
    assert publication.pages == "212-230"
    assert publication.book_title == (
        "Le jeu dans tous ses états. Approches pluridisciplinaires du phénomène ludique"
    )
    assert publication.publisher == "L’Harmattan"
    assert publication.publisher_city == "Paris"


def test_conference_publisher_defaults_and_joint_publishers() -> None:
    assert extract_publisher_metadata(
        "« Article », in Editors (éd.), Livre, EUD, 2025, p.1-10."
    ) == ("EUD", "Dijon")
    assert extract_publisher_metadata(
        "« Article », in Editors (éd.), Livre, Albolote (Granada), Comares, 2025."
    ) == ("Comares", "Albolote (Granada)")
    assert extract_publisher_metadata(
        "« Article », Ponta Delgada, Letras Levadas Edições et Fresno, "
        "Bruma Publications, 2025."
    ) == (
        "Letras Levadas Edições; Bruma Publications",
        "Ponta Delgada; Fresno",
    )
    assert extract_publisher_metadata(
        "« Article », Villeneuve-d’Asq, Temps, Mondes, Sociétés, 2026."
    ) == (None, "Villeneuve-d’Ascq")


def test_conference_container_drops_trailing_acts_description() -> None:
    publication = parse_citation(
        "« Avant-propos », in Le pouvoir du médecin au XIXe siècle, actes des "
        "journées d’études de Rouen et de Florence 2023-2024, en collaboration "
        "avec Michela Landi, Paris, Lettres modernes Minard, 2025, p.7-17.",
        "Communication dans un congrès",
        PublicationType.CONFERENCE_PAPER,
        1,
        "Florence Fix",
    )
    assert publication.book_title == "Le pouvoir du médecin au XIXe siècle"
    assert publication.conference_title == "Le pouvoir du médecin au XIXe siècle"


def test_parse_docx_sections_and_metadata(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Ouvrages")
    document.add_paragraph("Mon livre, Paris, Exemple, 2020, 230p.")
    document.add_paragraph("Article dans revue")
    document.add_paragraph("« Mon article », Revue test, 2024, p. 12-19. https://example.org/a")
    document.save(source)

    items = parse_docx(source, default_author="Florence Fix")

    assert len(items) == 2
    assert items[0].publication_type is PublicationType.BOOK
    assert items[0].year == 2020
    assert items[0].pages == "230"
    assert items[1].publication_type is PublicationType.JOURNAL_ARTICLE
    assert items[1].title == "Mon article"
    assert items[1].journal_title == "Revue test"
    assert items[1].url == "https://example.org/a"
    assert items[1].authors == ["Florence Fix"]


def test_leading_italic_span_is_used_as_complete_title(tmp_path: Path) -> None:
    source = tmp_path / "italic-title.docx"
    document = Document()
    document.add_paragraph("N°spécial de revue")
    paragraph = document.add_paragraph()
    title = "Growing Old in Nineteenth-Century France. Texts, Fictions, Representations"
    title_run = paragraph.add_run(title)
    title_run.italic = True
    paragraph.add_run(", numéro de L’Esprit créateur, Summer 2024, vol. 64, n°2, 132p.")
    document.save(source)

    items = parse_docx(source, default_author="Florence Fix")

    assert len(items) == 1
    assert items[0].title == title
    assert items[0].year == 2024
    assert items[0].pages == "132"


def test_nonitalic_space_inside_italic_title_is_preserved(tmp_path: Path) -> None:
    source = tmp_path / "century-title.docx"
    document = Document()
    document.add_paragraph("Ouvrages")
    paragraph = document.add_paragraph()
    first = paragraph.add_run("Femmes de spectacle au xixe")
    first.italic = True
    paragraph.add_run("\u00a0")
    century = paragraph.add_run("siècle")
    century.italic = True
    paragraph.add_run(", Paris, Exemple, 2022, 196p.")
    document.save(source)

    items = parse_docx(source, default_author="Florence Fix")

    assert items[0].title == "Femmes de spectacle au XIXe siècle"


def test_quoted_title_wins_over_later_italic_text(tmp_path: Path) -> None:
    source = tmp_path / "quoted-title.docx"
    document = Document()
    document.add_paragraph("Article dans revue")
    paragraph = document.add_paragraph()
    paragraph.add_run("« Mon article », in ")
    journal = paragraph.add_run("Revue test")
    journal.italic = True
    paragraph.add_run(", 2024, p.1-10.")
    document.save(source)

    items = parse_docx(source, default_author="Florence Fix")

    assert items[0].title == "Mon article"
    assert items[0].journal_title == "Revue test"


def test_url_only_paragraph_is_attached_to_previous_citation(tmp_path: Path) -> None:
    source = tmp_path / "continuation.docx"
    document = Document()
    document.add_paragraph("Communication dans un congrès")
    document.add_paragraph(
        "« Une communication », in Actes du colloque, Paris, 2025, p.19-31."
    )
    document.add_paragraph("https://example.org/book/99.")
    document.save(source)

    items = parse_docx(source, default_author="Florence Fix")

    assert len(items) == 1
    assert items[0].title == "Une communication"
    assert items[0].url == "https://example.org/book/99"
    assert "https://example.org/book/99" in items[0].raw_citation


def test_url_year_does_not_override_publication_year(tmp_path: Path) -> None:
    source = tmp_path / "url-year.docx"
    document = Document()
    document.add_paragraph("Article dans revue")
    document.add_paragraph(
        "« Article », Revue test, 2021, p.1-10. https://example.org/archive/2025"
    )
    document.save(source)

    items = parse_docx(source, default_author="Florence Fix")

    assert items[0].year == 2021


def test_generic_front_matter_uses_book_title_and_strips_author(tmp_path: Path) -> None:
    source = tmp_path / "front-matter.docx"
    document = Document()
    document.add_paragraph("Chapitre d’ouvrage")
    document.add_paragraph(
        "« Avant-propos », in Florence Fix, Tous malades. "
        "Représentations du corps souffrant, Paris, Orizons, 2018, p.9-21."
    )
    document.save(source)

    items = parse_docx(source, default_author="Florence Fix")

    assert items[0].book_title == "Tous malades. Représentations du corps souffrant"
    assert items[0].title == (
        "Avant-propos à Tous malades. Représentations du corps souffrant"
    )


def test_journal_issue_parts_are_kept_only_for_repeated_themes(tmp_path: Path) -> None:
    source = tmp_path / "issues.docx"
    document = Document()
    document.add_paragraph("N°spécial de revue")
    document.add_paragraph("Le Paon d’Héra 1, Orphée (1), 2006, 134p.")
    document.add_paragraph("Le Paon d’Héra 2, Orphée (2), 2007, 226p.")
    document.add_paragraph("Le Paon d’Héra 3, Roméo et Juliette (3), 2007, 275p.")
    document.add_paragraph("Le Paon d’Héra 11, Le roi Lear, 2018, 195p.")
    document.save(source)

    items = parse_docx(source, default_author="Florence Fix")

    assert [item.title for item in items] == [
        "Le Paon d’Héra 1, Orphée (1)",
        "Le Paon d’Héra 2, Orphée (2)",
        "Le Paon d’Héra 3, Roméo et Juliette",
        "Le Paon d’Héra 11, Le roi Lear",
    ]
    assert items[0].publisher == "Éditions du Murmure"
    assert items[0].publisher_city == "Neuilly-lès-Dijon"
    assert items[3].publisher == "Presses universitaires de Franche-Comté"
    assert items[3].publisher_city == "Besançon"


def test_generic_conference_front_matter_uses_proceedings_title(tmp_path: Path) -> None:
    source = tmp_path / "conference-front-matter.docx"
    document = Document()
    document.add_paragraph("Communication dans un congrès")
    document.add_paragraph(
        "« Avant-propos », avec Xavier Bonnier, in Xavier Bonnier et Florence Fix "
        "(éd.), Le Détour du comparant, Paris, Classiques Garnier, 2025, p.7-17."
    )
    document.save(source)

    items = parse_docx(source, default_author="Florence Fix")

    assert items[0].title == "Avant-propos au Détour du comparant"
    assert items[0].publisher == "Classiques Garnier"
    assert items[0].publisher_city == "Paris"


def test_generic_journal_titles_use_dossier_name(tmp_path: Path) -> None:
    source = tmp_path / "journal-front-matter.docx"
    document = Document()
    document.add_paragraph("Article dans revue")
    document.add_paragraph(
        "« Avant-propos », in Revue d’études culturelles, n°9, Jouer Marilyn, "
        "Dijon, ABELL, 2022, p.7-17."
    )
    document.add_paragraph(
        "« Introduction », Le Mélodramatique, revue Romantisme, n°208, "
        "2025/2, p.5-13."
    )
    document.add_paragraph(
        "« Introduction », en collaboration avec Thierry Roger, in Carnets du "
        "vivant, n°2 Flux et fleuves, 2025."
    )
    document.add_paragraph(
        "« Introduction », in Growing Old in Nineteenth-Century France. Texts, "
        "Fictions, Representations, numéro de L’Esprit créateur, Summer 2024, "
        "vol. 64, n°2, p.1-17."
    )
    document.add_paragraph(
        "« Avant-propos », Revue d’Études culturelles n°12 Fêtes de fin, "
        "Dijon, ABELL, 2025, p.7-12."
    )
    document.save(source)

    items = parse_docx(source, default_author="Florence Fix")

    assert [item.title for item in items] == [
        "Avant-propos à Jouer Marilyn",
        "Introduction au Mélodramatique",
        "Introduction à Flux et fleuves",
        "Introduction à Growing Old in Nineteenth-Century France. Texts, "
        "Fictions, Representations",
        "Avant-propos à Fêtes de fin",
    ]
